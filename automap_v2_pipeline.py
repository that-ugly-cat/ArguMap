"""AutoMap v2 — LLM extraction pipeline."""
from __future__ import annotations

import json
import re
from dataclasses import dataclass
from enum import Enum
from io import BytesIO
from pathlib import Path
from typing import List, Literal, Optional

import anthropic
from pydantic import BaseModel, model_validator

# ---------------------------------------------------------------------------
# Schema
# ---------------------------------------------------------------------------

class NodeType(str, Enum):
    CLAIM                   = "claim"
    NORMATIVE_PREMISE       = "normative_premise"
    EMPIRICAL_PREMISE       = "empirical_premise"
    METAPHYSICAL_COMMITMENT = "metaphysical_commitment"
    INTERMEDIATE_CONCLUSION = "intermediate_conclusion"
    OBJECTION               = "objection"


class StepAnnotation(BaseModel):
    valid:          Optional[bool] = None
    bias_label:     Optional[str]  = None
    bias_reason:    Optional[str]  = None
    fallacy_label:  Optional[str]  = None
    fallacy_reason: Optional[str]  = None


class ArgumentNode(BaseModel):
    id:      str
    type:    NodeType
    content: str
    notes:   Optional[str] = None


class InferentialStep(BaseModel):
    id:         str
    sources:    List[str]
    target:     str
    linked:     bool = False
    relation:   Literal["supports", "attacks", "qualifies"] = "supports"
    rule:        Optional[str] = None
    rule_reason: Optional[str] = None
    strength:    float = 0.5
    annotation:  StepAnnotation = StepAnnotation()


class ArgumentMapV2(BaseModel):
    id:    str
    title: str
    nodes: List[ArgumentNode]
    steps: List[InferentialStep]

    @model_validator(mode="after")
    def exactly_one_claim(self) -> ArgumentMapV2:
        claims = [n for n in self.nodes if n.type == NodeType.CLAIM]
        if len(claims) != 1:
            raise ValueError(f"Exactly one claim node required, found {len(claims)}")
        return self

    @property
    def claim(self) -> ArgumentNode:
        return next(n for n in self.nodes if n.type == NodeType.CLAIM)

    def node(self, node_id: str) -> ArgumentNode:
        return next(n for n in self.nodes if n.id == node_id)

    @classmethod
    def from_json(cls, path: str) -> ArgumentMapV2:
        with open(path, encoding="utf-8") as f:
            return cls.model_validate(json.load(f))

# ---------------------------------------------------------------------------
# Schemes
# ---------------------------------------------------------------------------

_SCHEMES_PATH = Path(__file__).parent / 'schemes.json'

with open(_SCHEMES_PATH, encoding='utf-8') as _f:
    SCHEMES = json.load(_f)


def _fmt_list(items: list) -> str:
    return '\n'.join(f'- {i["name"]}: {i["description"]}' for i in items)


def _flatten(section: str, key: str) -> list:
    out = []
    for family in SCHEMES[section].values():
        out.extend(family[key])
    return out


_RULES_REF     = _fmt_list(_flatten('inferential_rules', 'rules'))
_FALLACIES_REF = _fmt_list(_flatten('fallacies', 'fallacies'))
_BIASES_REF    = _fmt_list(_flatten('biases', 'biases'))

# ---------------------------------------------------------------------------
# Prompts
# ---------------------------------------------------------------------------

DEFAULT_MODEL = 'claude-sonnet-4-6'

SYSTEM = f"""\
You are an expert in analytic philosophy and argument mapping.
You extract structured ethical argument maps from academic texts.

## Node types
- claim                   — the central thesis being argued (exactly one per map)
- normative_premise       — a normative or ethical premise
- empirical_premise       — an empirical or factual premise
- metaphysical_commitment — a deep metaphysical assumption that grounds a normative premise
- intermediate_conclusion — a sub-conclusion within the inferential chain
- objection                — a counter-consideration raised against the claim, a premise, or an
  inference. Defined by its dialectical role (it tells against a position); connects via "attacks".

## ID conventions
- Claim:                    C1
- Normative premises:       N1, N2, ...
- Empirical premises:       E1, E2, ...
- Metaphysical commitments: M1, M2, ...
- Intermediate conclusions: IC1, IC2, ...
- Objections:               O1, O2, ...
- Inferential steps:        S1, S2, ...

## Notes field
The `notes` field has a specific meaning for each node type:

- **claim, normative_premise, empirical_premise, intermediate_conclusion, objection**: use `notes` for
  bibliographic references cited in the text in support of that node's content
  (e.g. "Ardila et al. 2019, Nature Medicine", "Helsinki Declaration Art. 26").
  If no reference is cited, leave `null`.

- **metaphysical_commitment**: use `notes` to identify the ethical or philosophical framework
  the commitment belongs to, when this can be established with reasonable confidence
  (e.g. "Kantian deontology", "principlism — beneficence/non-maleficence",
  "utilitarian consequentialism", "capabilities approach").
  Leave `null` if the commitment crosses frameworks or cannot be clearly attributed.

## Inferential rules (`rule` field on inferential steps)
Use the names below when possible. Use free text only if no entry fits.
{_RULES_REF}

## Fallacies (`fallacy_label` field inside `annotation` on inferential steps)
Use the names below when possible. Use free text only if no entry fits.
{_FALLACIES_REF}

## Biases (`bias_label` field inside `annotation` on inferential steps)
Use the names below when possible. Use free text only if no entry fits.
{_BIASES_REF}

## Output format
Each processing stage outputs a JSON block inside triple backticks tagged `json`.
Output ONLY valid JSON in that block — no trailing commas, no comments.

## Language
Detect the language of the input text and use that same language for ALL textual output: \
node content, claim text, notes, rule names (when using free text), fallacy/bias labels \
and reasons, rule reasons. Do not translate or switch languages.
"""

_STEP1_PROMPT = """\
## Task
Identify the central thesis (claim) of the following text.

Rules:
- There is exactly one claim per map.
- The claim is the proposition the author is ultimately arguing for — state the author's
  own thesis, whether it is affirmative or negative (e.g. "X should not be done" is valid).
- State it as a single, self-contained declarative sentence.
- Do NOT include supporting reasoning here.

## Text
{text}

## Output
First, reason briefly about what the central thesis is (2–4 sentences).
Write your reasoning in the same language as the text above.
Then output:

```json
{{"id": "C1", "type": "claim", "content": "...", "notes": null}}
```
"""

_STEP2_PROMPT = """\
## Context
The central claim has already been extracted:
{claim_json}

## Task
From the following text, extract all remaining argument nodes EXCEPT metaphysical commitments
(those come in the next step).

Extract:
- normative_premise  (N1, N2, ...)
- empirical_premise  (E1, E2, ...)
- intermediate_conclusion (IC1, IC2, ...) — sub-conclusions derived from premises that feed
  toward the main claim; extract them as nodes here; they will also appear as sources or
  targets in inferential steps (step 4)
- objection (O1, O2, ...) — a counter-consideration the text raises AGAINST the claim, a premise,
  or an inference (NOT a supporting premise). Objections are signalled by moves like "however",
  "critics argue", "one might object", "opponents claim", "a challenge is". Extract an objection
  even when the author then rebuts it — mapping the objection is precisely the point. In step 4
  an objection connects to whatever it challenges with relation "attacks".

Rules:
- Assign sequential IDs within each type (N1 before N2, etc.).
- Keep `content` self-contained and faithful to the text — do not paraphrase heavily.
- An objection is defined by its dialectical ROLE, not its content: it may be empirical or
  normative in substance, but if its function in the text is to tell against a position, type it
  `objection` — not `normative_premise`/`empirical_premise`. A consideration offered in favour of
  the claim is a premise; the same kind of consideration offered against it is an objection.
- If a premise mixes normative and empirical content, split it into two separate nodes:
  one normative_premise for the ethical or evaluative component, one empirical_premise
  for the factual component. Do not merge them into a single node.
- If the text cites a reference for a node, put it in `notes`.
  Associate references with the node they evidence, not the nearest node in the text.
- Omit nodes that are not actually used in the argument (background context only).

## Text
{text}

## Output
First, reason through which premises and intermediate conclusions are present (4–8 sentences).
Write your reasoning in the same language as the text above.
Then output:

```json
[{{"id": "N1", "type": "normative_premise", "content": "...", "notes": null}}, ...]
```
"""

_STEP3_PROMPT = """\
## Context
Nodes extracted so far:
{registry_json}

## Task
Identify the metaphysical commitments (M1, M2, ...) implicit in the normative premises above.

A metaphysical commitment is a deep background assumption about the nature of reality,
persons, agency, or value that a normative premise presupposes but does not argue for.
Examples: "persons are ends in themselves", "consciousness grounds moral status",
"technological tools are morally neutral absent governance".

Rules:
- Extract a metaphysical commitment only if making it explicit changes how one would evaluate
  the normative premise that presupposes it. If the commitment is irrelevant to assessing
  the premise's validity or force, skip it.
- Do NOT invent commitments not implicit in the text or the normative premises.
- Assign IDs M1, M2, ... in order.
- In `notes`, identify the ethical or philosophical framework the commitment belongs to
  (e.g. "Kantian deontology", "principlism — non-maleficence", "utilitarian consequentialism").
  Leave `null` if the commitment cannot be clearly attributed to a single framework.

## Text
{text}

## Output
First, reason about which normative premises have implicit metaphysical commitments (3–6 sentences).
Then output (empty array if none):

```json
[{{"id": "M1", "type": "metaphysical_commitment", "content": "...", "notes": "Kantian deontology"}}]
```
"""

_STEP4_PROMPT = """\
## Node registry (all available nodes)
{registry_json}

## Task
Map the inferential structure of the following text.
For each inferential move, produce one InferentialStep.

## InferentialStep schema
```json
{{
  "id":          "S1",
  "sources":     ["N1"],
  "target":      "C1",
  "linked":      false,
  "relation":    "supports",
  "rule":        "modus ponens",
  "rule_reason": null,
  "strength":    0.8,
  "annotation": {{
    "valid":          true,
    "bias_label":     null,
    "bias_reason":    null,
    "fallacy_label":  null,
    "fallacy_reason": null
  }}
}}
```

## Relations
- `supports`  — the sources make the target more credible.
- `attacks`   — the sources tell against the target. Every step whose source is an `objection`
  node is an `attacks` step; a step that rebuts an objection likewise `attacks` that objection.
- `qualifies` — the sources restrict the scope or conditions of the target.

## Annotation guide

Work through each inferential step in this order:

**1. Validity** — assess whether the inference holds given the stated premises:
- `true`  — the premises logically or evidentially support the target, even if not deductively certain
- `false` — the inference contains a recognisable logical flaw; always annotate `fallacy`
- `null`  — the text does not provide enough information to assess the inference

**2. Fallacy** — if `valid: false`, identify the fallacy.
- `fallacy_label`: name from the reference list, or free text if none fits.
- `fallacy_reason`: 1–2 sentences identifying which specific feature of this inferential step
  makes it an instance of that fallacy. Diagnose this step — do not define the fallacy.
  Concise and technical; presupposes the reader knows the label.
- Leave both `null` if `valid` is `true` or `null`.

**3. Bias** — independently of validity, flag distorting assumptions.
- `bias_label`: name from the reference list, or free text.
- `bias_reason`: 1–2 sentences identifying the specific distorting assumption operating in
  this step and how it skews the inference. Concise and technical; presupposes the reader
  knows the label.
- A step can have both fallacy and bias if the flaw is structural AND assumption-driven.
- Leave both `null` if the premises are transparent and the inference is well-grounded.

**4. Rule** — name the inferential rule the step instantiates.
- `rule`: name from the reference list, or free text.
- `rule_reason`: 1 sentence explaining how this step instantiates the rule.
  Include only when the fit is non-obvious. Leave `null` for straightforward cases.

**5. Strength** — evidential weight of the inference:
- 1.0 = direct logical entailment or definitional necessity
- 0.7–0.9 = strong empirical support, well-replicated evidence
- 0.4–0.6 = moderate support, plausible but contested
- 0.1–0.3 = weak or highly speculative
- 0.0 = nominal link only

## Structural rules
- Use ONLY node IDs present in the registry above. Do NOT invent new IDs.
- `linked: true` only when sources form a co-premise group (all required together).
- `relation`: `supports` strengthens the target, `attacks` undermines it, `qualifies` constrains it.

## Text
{text}

## Output
Output ONLY the JSON array of InferentialStep objects, no prose:

```json
[...]
```
"""

# ---------------------------------------------------------------------------
# Utilities
# ---------------------------------------------------------------------------

def _extract_json(text: str):
    m = re.search(r'```json\s*([\s\S]+?)```', text)
    if not m:
        raise ValueError("No ```json block found in response:\n" + text)
    return json.loads(m.group(1).strip())


def _reasoning_only(text: str) -> str:
    """Strip ```json ...``` blocks, return only the prose reasoning."""
    stripped = re.sub(r'```json[\s\S]*?```', '', text)
    return stripped.strip()


def _chat(system: str, user: str, *, client: anthropic.Anthropic,
          model: str, max_tokens: int = 8192) -> tuple[str, dict]:
    """Returns (content, usage) where usage = {'input_tokens': int, 'output_tokens': int}."""
    msg = client.messages.create(
        model=model,
        max_tokens=max_tokens,
        system=system,
        messages=[{"role": "user", "content": user}],
    )
    if msg.stop_reason == "max_tokens":
        raise ValueError(
            f"Response truncated (max_tokens={max_tokens}). "
            "Increase max_tokens or shorten the input text."
        )
    usage = {"input_tokens": msg.usage.input_tokens, "output_tokens": msg.usage.output_tokens}
    return msg.content[0].text, usage

# ---------------------------------------------------------------------------
# Pipeline result
# ---------------------------------------------------------------------------

@dataclass
class PipelineResult:
    argmap:       ArgumentMapV2
    claim:        dict
    nodes:        list
    metaphysical: list
    steps:        list
    raw:          dict   # {'claim': str, 'nodes': str, 'metaphysical': str, 'steps': str}
    usage:        dict   # {'input_tokens': int, 'output_tokens': int, 'model': str}

    def to_dict(self) -> dict:
        """Merge argmap + pipeline metadata into a single dict for export."""
        d = self.argmap.model_dump()
        d['_pipeline'] = {
            'claim':        self.claim,
            'nodes':        self.nodes,
            'metaphysical': self.metaphysical,
            'steps':        self.steps,
            'raw':          self.raw,
            'usage':        self.usage,
        }
        return d

    def save(self, path: str) -> None:
        """Write the merged dict to a JSON file."""
        Path(path).write_text(
            json.dumps(self.to_dict(), indent=2, ensure_ascii=False),
            encoding='utf-8',
        )

# ---------------------------------------------------------------------------
# Step functions
# ---------------------------------------------------------------------------

def step1_claim(text: str, *, client: anthropic.Anthropic,
                model: str = DEFAULT_MODEL) -> tuple[dict, str, dict]:
    raw, usage = _chat(SYSTEM, _STEP1_PROMPT.format(text=text),
                       client=client, model=model, max_tokens=512)
    return _extract_json(raw), raw, usage


def step2_nodes(text: str, claim: dict, *, client: anthropic.Anthropic,
                model: str = DEFAULT_MODEL) -> tuple[list, str, dict]:
    raw, usage = _chat(SYSTEM,
                       _STEP2_PROMPT.format(text=text,
                                            claim_json=json.dumps(claim, ensure_ascii=False)),
                       client=client, model=model, max_tokens=4096)
    return _extract_json(raw), raw, usage


def step3_metaphysical(text: str, registry: list, *, client: anthropic.Anthropic,
                       model: str = DEFAULT_MODEL) -> tuple[list, str, dict]:
    raw, usage = _chat(SYSTEM,
                       _STEP3_PROMPT.format(text=text,
                                            registry_json=json.dumps(registry, indent=2,
                                                                     ensure_ascii=False)),
                       client=client, model=model, max_tokens=3000)
    result = _extract_json(raw)
    return (result if isinstance(result, list) else []), raw, usage


def step4_steps(text: str, registry: list, *, client: anthropic.Anthropic,
                model: str = DEFAULT_MODEL) -> tuple[list, str, dict]:
    raw, usage = _chat(SYSTEM,
                       _STEP4_PROMPT.format(text=text,
                                            registry_json=json.dumps(registry, indent=2,
                                                                     ensure_ascii=False)),
                       client=client, model=model, max_tokens=8192)
    return _extract_json(raw), raw, usage

# ---------------------------------------------------------------------------
# Main pipeline
# ---------------------------------------------------------------------------

def extract_map(text: str, map_id: str, title: str,
                model: str = DEFAULT_MODEL,
                api_key: Optional[str] = None,
                on_step=None) -> PipelineResult:
    """
    Full pipeline: text → PipelineResult.

    on_step: optional callable(step_name: str, reasoning: str) called after each
             LLM step with the prose reasoning (JSON blocks stripped).
    """
    client = anthropic.Anthropic(api_key=api_key) if api_key else anthropic.Anthropic()

    def _notify(name, raw):
        if on_step is not None:
            on_step(name, _reasoning_only(raw))

    claim_node,  raw1, u1 = step1_claim(text, client=client, model=model)
    _notify('Step 1 — Claim', raw1)

    other_nodes, raw2, u2 = step2_nodes(text, claim_node, client=client, model=model)
    _notify('Step 2 — Nodes', raw2)

    registry = [claim_node] + other_nodes

    meta_nodes, raw3, u3 = step3_metaphysical(text, registry, client=client, model=model)
    _notify('Step 3 — Metaphysical commitments', raw3)
    registry = registry + meta_nodes

    steps, raw4, u4 = step4_steps(text, registry, client=client, model=model)
    _notify('Step 4 — Inferential steps', raw4)

    known_ids = {n["id"] for n in registry}
    for s in steps:
        bad = [sid for sid in s.get("sources", []) if sid not in known_ids]
        if s.get("target") not in known_ids:
            bad.append(s.get("target"))
        if bad:
            raise ValueError(f"Step {s['id']} references unknown IDs: {bad}")

    # Sanitize: a step's target cannot appear among its own sources (circular)
    for s in steps:
        sources = s.get("sources", [])
        target  = s.get("target")
        if target in sources:
            sources = [sid for sid in sources if sid != target]
            s["sources"] = sources
            s["linked"]  = len(sources) > 1

    total_in  = u1["input_tokens"]  + u2["input_tokens"]  + u3["input_tokens"]  + u4["input_tokens"]
    total_out = u1["output_tokens"] + u2["output_tokens"] + u3["output_tokens"] + u4["output_tokens"]

    argmap = ArgumentMapV2(id=map_id, title=title, nodes=registry, steps=steps)

    return PipelineResult(
        argmap=argmap,
        claim=claim_node,
        nodes=other_nodes,
        metaphysical=meta_nodes,
        steps=steps,
        raw={
            'claim':        raw1,
            'nodes':        raw2,
            'metaphysical': raw3,
            'steps':        raw4,
        },
        usage={
            'input_tokens':  total_in,
            'output_tokens': total_out,
            'model':         model,
        },
    )

# ---------------------------------------------------------------------------
# Ingest
# ---------------------------------------------------------------------------

def load_txt(path: str) -> str:
    return Path(path).read_text(encoding='utf-8')


def load_docx(path: str) -> str:
    from docx import Document
    doc = Document(path)
    return '\n'.join(p.text for p in doc.paragraphs if p.text.strip())


def normalize(text: str) -> str:
    text = re.sub(r'[ \t]+', ' ', text)
    text = re.sub(r'\n{3,}', '\n\n', text)
    return text.strip()


def ingest(path: str) -> str:
    """Load and normalize a .txt or .docx file from a filesystem path."""
    p = Path(path)
    suffix = p.suffix.lower()
    if suffix == '.txt':
        raw = load_txt(path)
    elif suffix in ('.docx', '.doc'):
        raw = load_docx(path)
    else:
        raise ValueError(f'Unsupported format: {suffix}')
    return normalize(raw)


def ingest_bytes(content: bytes, filename: str) -> str:
    """Load and normalize a .txt or .docx file from raw bytes (e.g. Streamlit file_uploader)."""
    suffix = Path(filename).suffix.lower()
    if suffix == '.txt':
        raw = content.decode('utf-8')
    elif suffix in ('.docx', '.doc'):
        from docx import Document
        doc = Document(BytesIO(content))
        raw = '\n'.join(p.text for p in doc.paragraphs if p.text.strip())
    else:
        raise ValueError(f'Unsupported format: {suffix}')
    return normalize(raw)
